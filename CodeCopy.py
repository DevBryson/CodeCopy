from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import os

def createDocment(title):
    global document

    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.size = Pt(9)
    document.add_heading(title, 0)


def copyCode(path, depth):
    global lineCount
    global fileCount

    files = os.listdir(path)
    for file in files:  # 遍历文件夹
        if not (os.path.isdir(path + "/" + file)):  # 判断是否是文件夹，不是文件夹才打开
            if not os.path.splitext(file)[1] == ".java" :       # 只拷贝java文件
                continue

            rootFileFp.write("|      " * depth + "+--" + file + "\n")
            # print("|      " * depth + "+--" + file)     #打印文件名
            fileCount += 1
            # 添加一个一级标题
            document.add_heading(file, level=1)

            fp = open(path + "/" + file, encoding="UTF-8")  # 打开文件
            iter_f = iter(fp);
            # 每读取一段写入一次
            paragraph = ""
            for line in iter_f:
                # print(line, end='')
                lineCount = lineCount + 1  # 统计行数
                if line == "\n":
                    document.add_paragraph(paragraph)
                    paragraph = ""
                else:
                    paragraph += line

            fp.close()
        else:
            rootFileFp.write("|      " * depth + "+--" + str(file) + "\n")
            # print("|      " * depth + "+--" + str(file))

            if file == "test" or file == "androidTest" or file == "build":      # 去除单元测试和生成的文件
                continue
            else:
                copyCode(path + "/" + file, depth + 1)


if __name__ == '__main__':
    # 初始化全局变量
    lineCount = 0
    fileCount = 0
    rootFileFp = open("工程结构.txt", mode='w', encoding="UTF-8")  # 打开文件

    print("请输入软著文档名：")
    title = input()
    createDocment(title=title)

    print("请输入源工程目录绝对路径：")
    path = input()

    print("正在拷贝中...")
    rootFileFp.write("项目工程目录绝对路径：" + path + "\n\n\n")
    copyCode(path=path, depth=0)

    # 保存软著文档
    document.save(title + '.docx')
    rootFileFp.close()

    print("代码拷贝成功！\n\n")
    print("软著文档路径：" + os.getcwd() + "\\" + title + '.docx')
    print("工程结构文件路径：" + os.getcwd() + "\\" + "工程结构.txt")
    print("总拷贝文件数：" + str(fileCount))
    print("总代码行数：" + str(lineCount))
